# Builds the Control Room documentation as a .docx.
# Run:  python build_docs.py <output.docx>

import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# The product palette, so the document reads as part of the same system.
INK        = RGBColor(0x0B, 0x0B, 0x0B)
SECONDARY  = RGBColor(0x52, 0x51, 0x4E)
MUTED      = RGBColor(0x89, 0x87, 0x81)
ACCENT     = RGBColor(0x2A, 0x78, 0xD6)
GOOD       = RGBColor(0x0C, 0xA3, 0x0C)
WARNING    = RGBColor(0xB0, 0x7A, 0x00)
CRITICAL   = RGBColor(0xD0, 0x3B, 0x3B)

SHADE_CODE = "F4F4F2"
SHADE_HEAD = "EDEDE9"
SHADE_WARN = "FDF6E3"
SHADE_CRIT = "FBEDED"
SHADE_INFO = "EEF4FC"

doc = Document()


# ----------------------------------------------------------------- plumbing --

def set_base_styles():
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.15
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), "Calibri")

    for name, size, color, before, after in [
        ("Heading 1", 20, ACCENT, 22, 8),
        ("Heading 2", 15, INK, 16, 6),
        ("Heading 3", 12, INK, 13, 4),
        ("Heading 4", 10.5, SECONDARY, 11, 3),
    ]:
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True


def shade(element, hexfill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexfill)
    element.append(shd)


CODE_COLOR = RGBColor(0x1C, 0x5C, 0xAB)


def add_runs(p, text, size=10.5, color=None, bold=False, italic=False):
    """Split on backticks so `like this` becomes inline code everywhere."""
    for j, chunk in enumerate(str(text).split("`")):
        if not chunk:
            continue
        r = p.add_run(chunk)
        if j % 2 == 1:
            r.font.name = "Consolas"
            r.font.size = Pt(size - 1)
            r.font.color.rgb = CODE_COLOR
        else:
            r.font.size = Pt(size)
            r.font.color.rgb = color or INK
            r.bold = bold
            r.italic = italic
    return p


def para(text="", bold=False, italic=False, size=10.5, color=None, space_after=7,
         align=None, style=None):
    p = doc.add_paragraph(style=style)
    add_runs(p, text, size=size, color=color, bold=bold, italic=italic)
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    return p


def rich(parts, space_after=7, style=None):
    """parts: list of (text, {bold,italic,code,color})"""
    p = doc.add_paragraph(style=style)
    for text, opts in parts:
        run = p.add_run(text)
        run.bold = opts.get("bold", False)
        run.italic = opts.get("italic", False)
        if opts.get("code"):
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = opts.get("color", RGBColor(0x1C, 0x5C, 0xAB))
        else:
            run.font.size = Pt(opts.get("size", 10.5))
            run.font.color.rgb = opts.get("color", INK)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(text, level=0, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.size = Pt(10.5)
    add_runs(p, text)
    return p


def numbered_list(items):
    """Numbers are written literally, so Word cannot continue a previous list
    and restart the count at 7."""
    for i, text in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.45)
        p.paragraph_format.first_line_indent = Inches(-0.28)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{i}.")
        r.bold = True
        r.font.size = Pt(10.5)
        r2 = p.add_run("\t")
        r2.font.size = Pt(10.5)
        add_runs(p, text)


def code_block(lines):
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.space_before = Pt(6 if i == 0 else 0)
        p.paragraph_format.space_after = Pt(6 if i == len(lines) - 1 else 0)
        p.paragraph_format.line_spacing = 1.0
        shade(p._p.get_or_add_pPr(), SHADE_CODE)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def callout(title, body, kind="info"):
    fill = {"info": SHADE_INFO, "warn": SHADE_WARN, "crit": SHADE_CRIT}[kind]
    bar = {"info": ACCENT, "warn": WARNING, "crit": CRITICAL}[kind]
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.cell(0, 0)
    shade(cell._tc.get_or_add_tcPr(), fill)
    cell.width = Inches(6.4)

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = bar

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    r2.font.size = Pt(10)
    r2.font.color.rgb = INK

    _cell_borders(cell, left=(bar, 18))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def _cell_borders(cell, left=None):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    if left:
        color, size = left
        el = OxmlElement("w:left")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:color"), "%02X%02X%02X" % (color[0], color[1], color[2]))
        borders.append(el)
    tcPr.append(borders)


def table(headers, rows, widths=None, font=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        shade(cell._tc.get_or_add_tcPr(), SHADE_HEAD)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(font)
        r.font.color.rgb = SECONDARY

    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            # `backticks` inside a cell render as inline code
            for j, chunk in enumerate(str(val).split("`")):
                if not chunk:
                    continue
                r = p.add_run(chunk)
                if j % 2 == 1:
                    r.font.name = "Consolas"
                    r.font.size = Pt(font - 0.5)
                    r.font.color.rgb = RGBColor(0x1C, 0x5C, 0xAB)
                else:
                    r.font.size = Pt(font)

    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def page_break():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_toc():
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-3" \h \z \u')
    fld.set(qn("w:dirty"), "true")
    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Right-click and choose Update Field to build the contents."
    run.append(t)
    fld.append(run)
    p._p.append(fld)


def update_fields_on_open():
    settings = doc.settings.element
    el = OxmlElement("w:updateFields")
    el.set(qn("w:val"), "true")
    settings.append(el)


def add_footer_page_numbers():
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("The Coordinators — Control Room documentation    ·    ")
    r.font.size = Pt(8)
    r.font.color.rgb = MUTED
    for instr in ["PAGE"]:
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), instr)
        run = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "16")
        rpr.append(sz)
        run.append(rpr)
        fld.append(run)
        p._p.append(fld)


set_base_styles()
for s in doc.sections:
    s.left_margin = s.right_margin = Inches(1.0)
    s.top_margin = Inches(0.9)
    s.bottom_margin = Inches(0.9)


# --------------------------------------------------------------- cover page --

for _ in range(5):
    doc.add_paragraph()

para("THE COORDINATORS", bold=True, size=11, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Control Room")
r.bold = True
r.font.size = Pt(38)
r.font.color.rgb = INK
p.paragraph_format.space_after = Pt(2)

para("Operating guide and technical documentation",
     size=13, color=SECONDARY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=28)

para("A web front door to the reporting playbook — the funnel end to end, "
     "plain-English questions answered from live systems, and a live view of "
     "what is actually answering.",
     size=11, color=SECONDARY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)

table(
    ["", ""],
    [
        ["Repository", "github.com/hilalaziz32/josh-operating-system"],
        ["Application", "public/ + api/ (Vercel, no build step, no dependencies)"],
        ["Reporting systems", "7 reporting APIs + Airtable (cross-system CPA)"],
        ["Model", "Gemini (gemini-3.6-flash by default), function calling"],
        ["Document version", "1.0 — 13 August 2026"],
        ["Audience", "Josh Klenoff (Parts 1–2); maintainers (Parts 3–6)"],
    ],
    widths=[1.6, 4.8],
    font=10,
)

page_break()

# ------------------------------------------------------------------- contents --

doc.add_heading("Contents", level=1)
para("This document has two halves. Parts 1 and 2 are for using the system and need no "
     "technical background. Parts 3 to 6 are for whoever maintains or deploys it.",
     color=SECONDARY, space_after=10)
add_toc()

page_break()

# =============================================================== PART 1 =====

doc.add_heading("Part 1 — Using the Control Room", level=1)

doc.add_heading("1.1  What this is", level=2)
para("The Coordinators run on eight separate systems. Each one holds a piece of the picture: "
     "what the ads cost, who applied, who is on the bench, who was screened, who was "
     "interviewed, who was assessed, and who was hired. Answering a simple question like "
     "“how are we doing?” used to mean opening eight dashboards and reconciling them by hand.")

para("The Control Room is a single page that does that for you. It has three views:")

bullet("shows the funnel end to end, pulled live from every system at once.", bold_lead="Dashboard ")
bullet("takes a question in plain English, works out which systems can answer it, "
       "asks them, and writes the report.", bold_lead="Ask ")
bullet("tells you which systems are actually answering right now, and what to do "
       "about any that are not.", bold_lead="Systems ")

para("It is read-only. Nothing in this application can change a record in any system. "
     "Every API key it holds is a reporting key, and the Airtable token is scoped to read only.",
     space_after=10)

callout(
    "The one rule that governs everything here",
    "Every number displayed came from exactly one system, and is labelled with that system. "
    "Nothing is averaged, reconciled, or recomputed. When two systems disagree, you are shown "
    "both figures and told they disagree — because the disagreement is usually the finding, "
    "and hiding it behind an average would destroy the only signal it carries.",
    "info",
)

doc.add_heading("1.2  Getting in", level=2)
para("Open the URL and enter the shared password. You stay signed in for twelve hours, "
     "then it asks again.")
para("If nobody can sign in at all, the deployment has no password set. That is deliberate: "
     "the application fails closed rather than putting live business data on an open URL. "
     "See Part 4 for how to set one.", space_after=10)

doc.add_heading("1.3  The Dashboard", level=2)
para("The Dashboard answers “what is the state of the business right now” without you asking "
     "anything. It loads every system in parallel and is usually complete within a few seconds.")

doc.add_heading("The funnel strip", level=3)
para("Eight tiles across the top, ordered as money in at the top and people out at the bottom:")
table(
    ["Stage", "Reported by", "What it counts"],
    [
        ["Spend", "Meta Ads", "Money spent on Meta ads in the window"],
        ["Leads", "Meta Ads", "Real leads from the CRM — never the Meta pixel count"],
        ["Applications", "Forms Platform", "Form submissions received"],
        ["Qualified", "Forms Platform", "Submissions marked qualified — the conversion event"],
        ["Screened", "Screening App", "Assessments completed"],
        ["Interviewed", "Video Interview", "AI-scored spoken interviews completed"],
        ["Assessed", "Candidate Assessment", "Proctored job simulations completed"],
        ["Hired", "HELM Ops (Signal)", "Candidates hired"],
    ],
    widths=[1.1, 1.7, 3.6],
)

callout(
    "The funnel is a shape, not a cohort",
    "Each stage is reported by a different system counting a different population. The 91 "
    "applications and the 2 screenings are not the same people at two moments — they are two "
    "systems measuring two different things. Read the strip as the shape of the funnel, never "
    "as one group of people followed through it.",
    "warn",
)

para("A stage showing a dash rather than a number means that system did not answer. "
     "That is not zero. “We do not know” and “nobody applied” are different facts and the "
     "Dashboard never conflates them.", space_after=10)

doc.add_heading("Caveats", level=3)
para("Below the funnel sit any caveats that apply to today's data. These are not generic "
     "warnings — each is a rule written against a known defect, and it only appears when the "
     "live data actually triggers it. Two exist today:")

table(
    ["Caveat", "Fires when", "Why it matters"],
    [
        ["Meta spend is $0 for this window",
         "Meta reports zero spend",
         "Cost per lead is spend divided by leads, so at zero spend it reads $0.00. That is "
         "meaningless, not excellent. It usually means the ad account is still disabled — and "
         "any zero further down the funnel may be that same cause, not a business result."],
        ["Two systems disagree on screenings",
         "The Screening App and HELM Ops report different screening counts",
         "They measure at different points in the process and are known to disagree. Both "
         "figures are shown as reported. Neither is adjusted to match the other."],
    ],
    widths=[1.7, 1.7, 3.0],
)

doc.add_heading("Needs attention", level=3)
para("Anything actionable, hoisted out of the detail and sorted by size: review backlogs, "
     "roles past their target date, ingest failures. If this panel is empty, nothing is "
     "flagged.", space_after=10)

doc.add_heading("System cards", level=3)
para("One card per system with its own metrics, each showing the value, the change against "
     "the previous period of equal length, and a note where a number needs one. "
     "A green dot and “answering” means the system responded. Where a metric is tagged "
     "“raw”, the label is the API's own field name — used when that system has not documented "
     "a friendlier one, so what you see stays traceable to what it sent.", space_after=10)

doc.add_heading("The window picker", level=3)
para("Last 7 days, Last 30 days, Last 90 days, or This month. The window applies to every "
     "system at once. Each system names its date parameter differently; the application "
     "translates for each one, so the whole page always covers the same period.", space_after=10)

doc.add_heading("1.4  Ask", level=2)
para("Type the question you were going to ask anyway. You never name a system, pick an "
     "endpoint or choose a date format.")

para("Six presets sit above the box for the questions asked most often. Click one and it runs. "
     "The window picker beside the Run button scopes the question; you can also just say it "
     "in words (“in June”, “last quarter”) and it will be honoured.")

doc.add_heading("What happens while it runs", level=3)
para("An activity log appears showing each system as it is consulted, with a green dot when "
     "it answers and a red one when it does not. The report then streams in as it is written. "
     "A single-system question typically takes 13–20 seconds; a full report across all seven, "
     "about 30 seconds.")

para("When it finishes you can copy the report to the clipboard or download it as a Markdown "
     "file — useful for pasting into a board update or an email.", space_after=10)

doc.add_heading("Typical questions", level=3)
table(
    ["Ask", "What happens"],
    [
        ["Give me this week's report", "All seven systems, merged in funnel order, with a "
         "cross-system summary on top"],
        ["What needs my attention?", "All seven, filtered to only the bullets that name a "
         "problem — a backlog, a failure, a slipping role, a rising cost"],
        ["How's Meta looking?", "Meta Ads alone"],
        ["Are we wasting money on any campaigns?", "Meta Ads, drilled into campaigns and "
         "creatives — zero-lead campaigns and cost per lead by ad"],
        ["How's screening looking?", "All three assessment systems, which are separate "
         "applications, plus what they say together"],
        ["Which sourcing channel converts best?", "Screening App — volume and pass rate per "
         "channel and per sourcer"],
        ["Who's waiting on a decision from me?", "The review queues across every system, as "
         "counts, with where to go and clear each one"],
        ["Which roles are stalling?", "HELM Ops — roles past target date, roles with no "
         "candidates, and the staffing forecast"],
        ["What's my CPA?", "Meta spend joined to confirmed orders in Airtable. All-time by "
         "design, and never bundled into a weekly report"],
        ["Full picture for June", "All seven, scoped to 1–30 June"],
    ],
    widths=[2.2, 4.2],
)

doc.add_heading("Bookmarking a question", level=3)
para("A question can travel in the URL, so a recurring ask becomes a link you can save:")
code_block(["https://<your-app>/?q=What%20needs%20my%20attention"])
para("Open the bookmark and it signs you in and runs straight away.", space_after=10)

doc.add_heading("1.5  Systems", level=2)
para("A live probe of every reporting API — not a check that a key exists, but an actual "
     "request. It distinguishes states that look identical from the outside:")

table(
    ["Status", "Meaning", "What to do"],
    [
        ["answering", "Responded normally, with the round-trip time", "Nothing"],
        ["not configured", "This deployment has no URL or key for it", "Set the named "
         "variables in the Vercel project"],
        ["key rejected", "The key was sent and refused", "Two ends must match — check "
         "`REPORTING_API_KEY` in that system's own project"],
        ["erroring", "The host answered, but with an error", "The host's own message is "
         "shown; act on it"],
        ["timed out / unreachable", "No response at all", "Check the host is still deployed"],
    ],
    widths=[1.3, 2.1, 3.0],
)

para("A count in the left-hand rail shows how many are answering, so a problem is visible "
     "from any view without opening this one.", space_after=10)

doc.add_heading("1.6  How to read the numbers honestly", level=2)
para("These are the traps that make a confident report wrong. The application encodes most "
     "of them, but they are worth knowing.")

doc.add_heading("Unavailable is not zero", level=3)
para("A system that fails to answer shows a dash and is named in a “not available this run” "
     "line. It never contributes a zero to anything.")

doc.add_heading("A $0 cost per lead is not good news", level=3)
para("Cost per lead is spend divided by CRM leads. With the Meta ad account disabled, spend "
     "is $0 and so is the cost per lead. The Dashboard flags this.")

doc.add_heading("Fit and Pass are two different bars", level=3)
para("In Candidate Assessment, the platform's “Fit” threshold is 7.0/10 while the customer's "
     "own rubric passes at 40/50, which is 8.0/10. A candidate can be Fit and Borderline at "
     "the same time. Fit rate is reported by default and the two are never merged into a "
     "single “pass rate”.")

doc.add_heading("Some scores are a floor, not a verdict", level=3)
para("About 9 of the 50 rubric points in Candidate Assessment can only be awarded by a human "
     "watching the screen recording, and count as zero until that happens. An average quoted "
     "before review is therefore a floor. Reports quote the pending count alongside.")

doc.add_heading("The AI grader currently reads harsher than people", level=3)
para("As of 13 August 2026, reviewers agreed with the Candidate Assessment AI on 40% of "
     "submissions, overrode 9 of 15, and revised upward in 8 of those 9. The mean gap is "
     "7.14 points out of 50 and is widening. Treat an AI score as evidence, never a verdict, "
     "and read it next to the override rate. The Video Interview AI is far better calibrated "
     "— about 89% pass/fail agreement.")

doc.add_heading("Forms Platform volume is inflated by test data", level=3)
para("A form named “Test Form” and a job named “Dummy Job” account for a large share of "
     "submissions — in one recent week, 52 of 91, with the dummy job qualifying at 100%. "
     "Headline growth percentages from this system should be read with that in mind until "
     "the test records are removed.")

doc.add_heading("Job attribution on the Forms Platform is broken", level=3)
para("Roughly half of applications carry job_id 347, a legacy default, because most traffic "
     "still lands on an older form with no role picker. Role demand from this system is "
     "therefore partial.", space_after=10)

doc.add_heading("1.7  What it will not do", level=2)
para("This is a reporting surface. It answers how many, never who.")
bullet("It does not display candidate names, email addresses, phone numbers, or free-text answers.")
bullet("Asking for them directly returns a refusal and the equivalent count instead.")
bullet("Per-candidate endpoints are blocked in the server, not merely discouraged in an instruction.")
para("“Who is waiting on a decision?” is answered as a queue: how many, in which system, for "
     "which job, and where to go and clear it. If you need an individual's details, open that "
     "system's own admin interface, where the access is logged and appropriate.", space_after=6)

page_break()

# =============================================================== PART 2 =====

doc.add_heading("Part 2 — The systems", level=1)
para("Seven reporting systems plus Airtable. Each is independent, with its own deployment, "
     "its own key, and its own idea of a date range.", space_after=10)

SYSTEMS = [
    ("Meta Ads", "Meta Stats",
     "Paid acquisition: spend, real cost per lead, campaign and creative performance, "
     "wasted spend, and alerts.",
     [("Leads means CRM leads", "This account has no Meta CAPI, so the pixel systematically "
       "under-counts. The real number comes from the GoHighLevel CRM, attributed back by UTM. "
       "Pixel counts are never reported as leads."),
      ("Alerts always cover 7 days", "Regardless of the window asked for. By design in that system."),
      ("The ad account was disabled on 2 July 2026", "Spend went to zero and the whole funnel "
       "below it went quiet. Re-check this before reading any zero as a business result.")]),

    ("Forms Platform", "forms.coordinators.pro",
     "Application volume, qualified conversion, which roles attract applicants, the "
     "unreviewed backlog, contactability, and testimonials.",
     [("Qualified is the money metric", "Marking a submission qualified fires its automation, "
       "so it is the closest thing this system has to “good candidate found”."),
      ("Status lifecycle", "pending → reviewed → qualified → published → archived."),
      ("Test data inflates the totals", "See Part 1.6.")]),

    ("Candidate Inventory", "Candidate Inventory",
     "Offshore, remote-only healthcare candidate supply: bench size, intake volume and trend, "
     "country and skill concentration, pipeline stages, and ingest failures.",
     [("Two inventories", "future_potential is the curated, placeable bench. all is every "
       "stage. Always be explicit about which is being reported."),
      ("Not every failure is a fault", "NOT_REMOTE_CAPABLE rejections are the remote-only "
       "business rule working as intended. PARSE_FAIL, API_ERROR and VALIDATION_FAIL are real "
       "problems."),
      ("Totals versus window", "Bench totals are all-time; window figures are what happened "
       "inside the period. They are never conflated.")]),

    ("Screening App (Assessments)", "screening.coordinators.pro",
     "Intelligence, Big Five personality, verbal and writing assessments; recruiter pass and "
     "fail; channel and sourcer performance; anti-cheat flags.",
     [("Pass rate covers decided attempts only", "passed divided by passed plus failed, so a "
       "growing review backlog does not artificially depress it."),
      ("Awaiting review is an action item", "Not a quality problem — it is work waiting on a "
       "human."),
      ("Disqualified is not failed", "The attempt was voided, not judged poor. It is never "
       "folded into the pass rate.")]),

    ("Video Interview", "interview.coordinators.pro",
     "Writing task, Big Five, and an AI-scored spoken interview; AI pass/fail recommendations; "
     "AI-versus-reviewer calibration; anti-cheat.",
     [("A separate application", "Despite the similar name, this is not the Screening App. "
       "Tell them apart by their endpoints, never by name."),
      ("Awaiting review is all-time on purpose", "A candidate waiting three weeks is still "
       "waiting, whatever window you asked for."),
      ("Calibration needs scored pairs", "If nobody has personally scored a candidate, the "
       "AI's accuracy is unknown — which is not the same as good.")]),

    ("Candidate Assessment", "assesment.coordinators.pro",
     "A proctored 30-minute job simulation with webcam, screen recording and a mock patient "
     "call, scored by AI against a competency rubric.",
     [("Live since 13 August 2026", "Any window reaching back earlier under-reports it. "
       "A thin section for an older window is history, not a fault."),
      ("Opened is not applied", "The funnel starts when an invited candidate opens their "
       "private link, so the first drop-off is invitees who never began."),
      ("Two pass lines, and pending points", "See Part 1.6.")]),

    ("HELM Ops (Signal)", "Airtable via the Signal app",
     "The recruiting funnel from Airtable: screenings, candidate flow and hires, supply and "
     "demand spend, open roles, time to placement, and a deterministic staffing forecast.",
     [("No model sits between the data and the number", "Every figure is computed "
       "deterministically, so they are auditable and stable."),
      ("Read the confidence", "This base holds only a few months of history, so most rates "
       "come back Low. A Low-confidence figure is a rough guide, never a forecast."),
      ("It redacts people server-side", "Names, emails, phones, attachments and notes never "
       "leave that system, which is why its candidate endpoint stays available here.")]),
]

for name, host, covers, notes in SYSTEMS:
    doc.add_heading(name, level=2)
    rich([("Source: ", {"bold": True, "color": SECONDARY, "size": 10}),
          (host, {"code": True})], space_after=4)
    para(covers, space_after=6)
    for title, body in notes:
        rich([(title + " — ", {"bold": True}), (body, {})], space_after=3)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

doc.add_heading("CPA (cross-system)", level=2)
para("Not a reporting API. CPA is the answer to “what does a paying client actually cost”, "
     "and it spans three systems:")
code_block([
    "Meta ad spend  ->  GHL CRM lead  ->  Airtable Lead  ->  Airtable Role  =  a confirmed order",
])
para("A confirmed order is a Roles record whose Lead link is set — not the Ordered checkbox, "
     "which disagrees with it, and not Order Confirmed?, which is ticked on nothing. "
     "CPA is spend divided by the count of roles linked to a Meta-sourced lead.")
para("It is all-time by design, because the order side cannot be reliably date-windowed, and "
     "it is deliberately excluded from the weekly report bundle. Return on ad spend cannot be "
     "computed at all today: Deal Value is empty on every ordered lead, so we can say what an "
     "order costs but not what it is worth. The system refuses to print a fabricated ROAS.",
     space_after=10)

callout(
    "Attribution rests on a hand-set label",
    "A lead counts as Meta when its Source field says Meta. That is typed by a person, not a "
    "click ID. Meta CAPI is largely absent, so CPA is a close operational estimate rather than "
    "a fact. Wiring CAPI properly remains the highest-leverage fix available to this stack.",
    "warn",
)

page_break()

# =============================================================== PART 3 =====

doc.add_heading("Part 3 — How it works", level=1)

doc.add_heading("3.1  Shape of the thing", level=2)
para("Two front doors onto one playbook:")
table(
    ["Surface", "For", "What it is"],
    [
        ["Control Room", "Josh", "This web application, deployed on Vercel"],
        ["Claude Code", "Maintainers", "Open the repository and ask; the same skills, plus "
         "everything else the tool can do"],
    ],
    widths=[1.4, 1.2, 3.8],
)

para("The playbook itself lives in `.claude/skills/`. Each system has a SKILL.md describing "
     "its endpoints, its date parameters, how to read its fields, and what its output must "
     "look like. A master skill sits on top and does routing and merging. The web application "
     "does not duplicate any of that — it reads the same files at runtime, so there is one "
     "source of truth.")

doc.add_heading("3.2  File map", level=2)
table(
    ["Path", "Responsibility"],
    [
        ["`lib/systems.js`", "The catalog. One entry per system: environment variables, auth "
         "style, probe endpoint, how a window maps to that system's parameters, and which "
         "response fields are safe to put on a tile."],
        ["`lib/upstream.js`", "The only code that holds a key or calls an upstream. Timeouts, "
         "the shared response envelope, PII enforcement, tile extraction."],
        ["`lib/auth.js`", "The password gate: HMAC-signed, HTTP-only session cookie."],
        ["`lib/http.js`", "Request and response helpers shared by the routes."],
        ["`api/session.js`", "Sign in, sign out, and session state."],
        ["`api/systems.js`", "Live probe of every system."],
        ["`api/dashboard.js`", "Parallel sweep, tile extraction, funnel assembly, caveat rules."],
        ["`api/ask.js`", "The Gemini function-calling loop, streamed over Server-Sent Events."],
        ["`public/`", "The client. No framework, no build step."],
        ["`public/markdown.js`", "The report renderer, isolated so it can be tested outside a "
         "browser."],
        ["`dev.js`", "A local server that reproduces Vercel's routing, so the app runs without "
         "the Vercel CLI."],
    ],
    widths=[1.7, 4.7],
)

doc.add_heading("3.3  The Dashboard request", level=2)
numbered_list([
    "The browser asks for a window.",
    "The server resolves every configured system and calls each one's summary endpoint in "
    "parallel, each with its own 15-second timeout.",
    "Each response is turned into tiles using paths declared in the catalog. Those paths were "
    "transcribed from each system's SKILL.md, not guessed — where a field name was never "
    "documented, tiles are built from the response's own keys and marked “raw”.",
    "The funnel is assembled by pulling one declared value from each system.",
    "The caveat rules run against the live data.",
    "Raw payloads are discarded before the response is sent.",
])

callout(
    "Absence is preserved, never coerced",
    "If a declared field is missing, the tile is omitted rather than shown as zero. If a system "
    "fails, its funnel stage is blank rather than zero. This is the single most important "
    "property of the data path: the application would rather show you less than show you "
    "something untrue.",
    "info",
)

doc.add_heading("3.4  The Ask loop", level=2)
para("A question is answered by a model that has the playbook as its brief and two tools.")

table(
    ["Tool", "What it does"],
    [
        ["`read_playbook(system)`", "Returns that system's SKILL.md — its endpoints, date "
         "parameters, and reading rules."],
        ["`query_system(system, path, params)`", "Performs one read-only GET against that "
         "system's reporting API and returns the JSON. The server resolves the base URL and "
         "the key; the model only ever names a system."],
    ],
    widths=[2.3, 4.1],
)

para("The system instruction is assembled at request time from MEMORY.md and the master "
     "skill — roughly 28,000 characters — plus a roster of which systems are currently "
     "configured. The loop runs until the model stops calling tools, capped at 12 rounds and "
     "40 tool calls. Tool results are capped at 24,000 characters each.")

para("Routing comes first, from the catalog already in the instruction. Only then are the "
     "playbooks of systems it has decided to query read. This ordering matters: before it was "
     "enforced, a single-system question could read six playbooks and take twice as long.")

doc.add_heading("Two details specific to Gemini", level=3)
rich([("Frame separators. ", {"bold": True}),
      ("Gemini delimits its event stream with CRLFCRLF, not LFLF. A parser that matches only "
       "on a blank line finds no boundary at all, and the entire response backs up unread. "
       "The parser matches both, and flushes whatever remains when the stream ends, because "
       "the final frame is not always terminated.", {})])
rich([("Thought signatures. ", {"bold": True}),
      ("Gemini 3 attaches an encrypted handle on its reasoning to each part. The API is "
       "stateless, so anything dropped when the turn is echoed back cannot be recovered. Parts "
       "are therefore preserved verbatim; only adjacent unsigned plain text is merged.", {})],
     space_after=10)

doc.add_heading("3.5  Authentication", level=2)
para("One shared password in an environment variable buys an HMAC-signed, HTTP-only cookie "
     "valid for twelve hours. The cookie carries only an expiry and its signature — nothing "
     "secret, and nothing forgeable without the signing key. Every API route refuses without it.")
para("With no password configured, nobody can sign in. That is the safe failure, not a bug.")

callout(
    "The password is the only real defence",
    "The login rate limiter counts attempts in memory, and a serverless platform creates fresh "
    "instances constantly, so the count resets out from under itself. It deters a lazy script "
    "and nothing more. Use a long passphrase, and consider enabling Vercel's own Deployment "
    "Protection as a second layer.",
    "crit",
)

doc.add_heading("3.6  Privacy enforcement", level=2)
para("The privacy rule is enforced in the server, in three layers, so it is a property of the "
     "system rather than an instruction a model is trusted to follow.")
table(
    ["Layer", "What it does"],
    [
        ["Flag stripping", "Parameters that would lift redaction — `include_pii`, "
         "`include_contact`, `unredacted` and similar — are removed from every outgoing "
         "request."],
        ["Person-level block", "Paths that return one row per person — `candidates`, "
         "`submissions`, and their by-id forms — are refused before any request is made. "
         "HELM Ops is exempt, declared in the catalog, because it redacts server-side."],
        ["Airtable scrubbing", "Airtable has no reporting API and returns whole rows, so "
         "fields whose names look personal are stripped from the payload before it reaches the "
         "model, and the withheld field names are listed so a thin response is never mistaken "
         "for missing data."],
    ],
    widths=[1.5, 4.9],
)
para("No reporting ability is lost. Every awaiting-review count already comes from an "
     "aggregate endpoint.", space_after=10)

doc.add_heading("3.7  Rendering", level=2)
para("Reports arrive as Markdown and are rendered by `public/markdown.js`, which escapes "
     "first and formats second, so no value from an upstream system can smuggle markup into "
     "the page. It is deliberately small — it covers headings, bullets including nested ones, "
     "tables, emphasis, inline code and rules, and nothing else.", space_after=6)

page_break()

# =============================================================== PART 4 =====

doc.add_heading("Part 4 — Deploying and running it", level=1)

doc.add_heading("4.1  Deploying to Vercel", level=2)
numbered_list([
    "Import the GitHub repository as a new Vercel project.",
    "Leave the preset as Other and the Root Directory as `./`. Set no build command — "
    "`vercel.json` already pins the framework as none and the output directory as `public`.",
    "Add the environment variables below. The quickest route is Import .env against your local "
    "file, then change `APP_PASSWORD` to a production-only value.",
    "Deploy, then open Systems. Eight of eight answering means everything is wired.",
])

para("There are no dependencies and no build step. Nothing is installed at deploy time.",
     space_after=10)

doc.add_heading("4.2  Environment variables", level=2)
table(
    ["Variable", "Required", "Purpose"],
    [
        ["`APP_PASSWORD`", "Yes", "The gate. Unset means nobody can sign in."],
        ["`META_STATS_API_URL` / `_KEY`", "Yes", "Meta Ads"],
        ["`FORMS_REPORTING_BASE_URL` / `_API_KEY`", "Yes", "Forms Platform"],
        ["`CANDIDATE_INVENTORY_API_URL` / `_KEY`", "Yes", "Candidate Inventory"],
        ["`SCREENING_API_BASE_URL` / `SCREENING_API_KEY`", "Yes", "Screening App (Assessments)"],
        ["`SCREENING_APP_URL` / `_API_KEY`", "Yes", "Video Interview"],
        ["`ASSESSMENT_REPORTING_URL` / `_KEY`", "Yes", "Candidate Assessment"],
        ["`HELM_OPS_BASE_URL` / `_API_KEY`", "Yes", "HELM Ops (Signal)"],
        ["`AIRTABLE_API_KEY` / `AIRTABLE_BASE_ID`", "CPA only", "The Meta-to-orders join"],
        ["`GEMINI_API_KEY`", "Optional", "Switches the Ask console on. Without it, Dashboard "
         "and Systems work unchanged and Ask reports itself off."],
        ["`GEMINI_MODEL`", "Optional", "Defaults to gemini-3.6-flash. Must support function "
         "calling."],
        ["`SESSION_SECRET`", "Optional", "Signs the cookie. Derived from APP_PASSWORD when "
         "unset, which means changing the password signs everyone out."],
    ],
    widths=[2.5, 0.9, 3.0],
)
para("Set them for Production and Preview. `.env` is gitignored and never deploys — Vercel "
     "reads its own environment.", space_after=10)

doc.add_heading("4.3  Running locally", level=2)
code_block([
    "cd \"josh - operating system\"",
    "npm run dev            # http://localhost:4321",
])
para("The local server reads `.env` directly and reproduces Vercel's routing. Route files are "
     "re-imported per request, so editing one needs no restart; changes to `lib/` or `.env` do. "
     "There is no install step.", space_after=10)

doc.add_heading("4.4  Timing and limits", level=2)
table(
    ["Question type", "Typical", "Tool calls"],
    [
        ["Single system", "13–20 seconds", "2–4"],
        ["Full weekly report", "about 30 seconds", "15–18"],
        ["Full report over an explicit date range", "about 40 seconds", "about 21"],
    ],
    widths=[2.8, 1.8, 1.8],
)
para("`api/ask.js` requests a 300-second ceiling, which plans with Fluid compute honour. "
     "The Hobby tier caps at 60 seconds, which today's full report fits inside — but without "
     "much headroom if an upstream is slow. If a report is cut short, the interface says so "
     "rather than presenting half a report as whole.", space_after=10)

doc.add_heading("4.5  Troubleshooting", level=2)
table(
    ["Symptom", "Likely cause", "Fix"],
    [
        ["Nobody can sign in", "No APP_PASSWORD on the deployment", "Set it and redeploy"],
        ["Ask tab says it is switched off", "No GEMINI_API_KEY", "Add it and redeploy"],
        ["A system reads “not configured”", "Missing URL or key", "The Systems view names the "
         "exact variables"],
        ["A system reads “key rejected”", "Key mismatch between the two ends", "Check "
         "`REPORTING_API_KEY` in that system's own project"],
        ["A system reads “erroring”", "The host answered with a fault", "Its own message is "
         "shown in the Systems view"],
        ["A report stops early", "Function time limit reached", "Ask a narrower question, or "
         "move to a plan with a higher ceiling"],
        ["Port 4321 in use locally", "A previous server is still running", "Stop it, or run "
         "`PORT=4322 npm run dev`"],
    ],
    widths=[1.9, 2.1, 2.4],
)

page_break()

# =============================================================== PART 5 =====

doc.add_heading("Part 5 — Maintaining it", level=1)

doc.add_heading("5.1  MEMORY.md is the ground truth", level=2)
para("MEMORY.md at the repository root holds what has been learned the hard way: what a field "
     "actually means, which numbers are known to be soft, which zeros are artefacts. The Ask "
     "console loads it into every request, and the Dashboard's caveat rules are written "
     "against it.")
para("It holds two kinds of entry. Semantics are stable and can be trusted. Known defects "
     "should be verified before being relied upon — and when one is fixed, the entry is "
     "deleted rather than left to rot, because a stale warning makes people distrust good data.")
para("When a defect is retired, remove its caveat rule from `api/dashboard.js` at the same "
     "time.", space_after=10)

doc.add_heading("5.2  Adding a system", level=2)
numbered_list([
    "Drop its skill folder into `.claude/skills/<name>/` with its SKILL.md.",
    "Add a row to the catalog table in the master skill, including the words Josh would use "
    "when he wants it — that table is what the routing reads.",
    "Slot it into the funnel order in the same file's merge section.",
    "Add an entry to `SYSTEMS` in `lib/systems.js`: its variables, auth style, probe endpoint, "
    "window mapping, and tile paths taken from its own SKILL.md.",
    "Add its variables to `.env.example` and to the preflight script.",
])
para("Nothing else changes. The new system owns its own API calls, its own date mapping and "
     "its own output; the application only needs to know it exists.", space_after=10)

doc.add_heading("5.3  The contract every system keeps", level=2)
bullet("Emit a `### <System Name>` block: insight bullets, then a data-window footer.")
bullet("Bullets are insights — what changed, what is trending, what needs attention — never a "
       "dump of records.")
bullet("Accept a date window and cover exactly that window.")
bullet("When unreachable, unconfigured or erroring, emit exactly “- data unavailable”. "
       "Never error out, never guess a number.")
para("")

doc.add_heading("5.4  Known defects at time of writing", level=2)
para("Carried from MEMORY.md as of 13 August 2026. Verify before relying on any of them.")
table(
    ["Defect", "Effect"],
    [
        ["Unauthenticated candidate media in the Candidate Assessment app",
         "Per-question answer videos are served without a session check. That application is "
         "now in production, so this is live. It is in that repository, not this one. "
         "Highest priority in the stack."],
        ["Meta CAPI largely absent",
         "Attribution rests on a hand-set Source label, so CPA is an estimate."],
        ["Deal Value empty on ordered leads",
         "Return on ad spend cannot be computed at all."],
        ["Ordered checkbox disagrees with the Roles link",
         "Order counts differ depending on which field is trusted. The link is the real one."],
        ["12 of 38 roles have no lead link",
         "Unattributed orders. If any are Meta's, true CPA is better than reported."],
        ["Lead and order counts disagree across systems",
         "Airtable, the CRM and HELM Ops give materially different numbers. Report both, name "
         "both, never average."],
        ["Screening volume disagrees between systems",
         "Surfaced automatically as a Dashboard caveat."],
        ["Forms Platform job attribution broken",
         "Most submissions carry a legacy job id, so role demand is partial."],
        ["Test records in the Forms Platform",
         "A test form and a dummy job inflate volume and qualification rate."],
    ],
    widths=[2.4, 4.0],
)

page_break()

# =============================================================== PART 6 =====

doc.add_heading("Part 6 — What has been verified", level=1)
para("The application was exercised against the live systems on 13 August 2026, using the "
     "questions Josh would actually ask. This section records what was checked, so a future "
     "reader knows what is evidence and what is assumption.")

doc.add_heading("6.1  Coverage", level=2)
table(
    ["Area", "Result"],
    [
        ["All eight systems probed", "8 of 8 answering"],
        ["Dashboard across all four windows", "7 of 7 systems, funnel complete in each"],
        ["Every system asked individually", "Correct routing and correct block format"],
        ["CPA cross-system join", "Recomputed live from Airtable; $1,314.58 per order; "
         "correctly refused to print a ROAS"],
        ["Multi-system routing", "“How's screening looking” spanned all three assessment "
         "systems and produced a genuine cross-system insight"],
        ["Uncovered topic", "Declined plainly, without stretching a system to cover it"],
        ["Explicit date range", "June scoped correctly across all systems"],
        ["Direct request for names and contact details", "Refused; answered as queue counts"],
        ["Report rendering", "Structural checks pass; upstream markup is escaped, not executed"],
        ["Browser pass", "Sign-in, Dashboard, Ask streaming and Systems all verified visually"],
    ],
    widths=[2.4, 4.0],
)

doc.add_heading("6.2  Defects found and fixed during testing", level=2)
table(
    ["Found", "Resolution"],
    [
        ["The Ask console returned empty reports against the real API",
         "Two causes: the event-stream frame separator, and dropped thought signatures. Both "
         "fixed; see 3.4."],
        ["Candidate names were printed in answer to “who is waiting on a decision”",
         "Person-level endpoints blocked server-side and Airtable rows scrubbed. The answer is "
         "now a queue count with where to clear it."],
        ["A single-system question read six playbooks",
         "Routing now happens from the catalog before any playbook is read. That question went "
         "from 11 tool calls and 29 seconds to 3 and 16."],
        ["Reports restated themselves",
         "One system in the answer now means exactly one block."],
        ["A relabelled section's footer disagreed with its heading",
         "Footers must now name the system exactly as the heading does."],
        ["The synthesis heading said “this week” on monthly windows",
         "Corrected in the master playbook, so the terminal benefits too."],
        ["Dashboard tile separators broke on rows of five",
         "Rebuilt using grid gaps, correct at any tile count."],
        ["A port clash printed an unhandled stack trace",
         "The local server now explains the clash and suggests another port."],
    ],
    widths=[2.7, 3.7],
)

doc.add_heading("6.3  What has not been verified", level=2)
bullet("The Airtable-side defects in MEMORY.md (items 1 to 7) were last checked on 13 July "
       "2026 and have not been re-verified since.")
bullet("Behaviour under a genuinely slow or failing upstream has been reasoned about and "
       "coded for, but not induced deliberately.")
bullet("Concurrent use by several people at once has not been load-tested.")

para("")
para("— End of document —", italic=True, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)

add_footer_page_numbers()
update_fields_on_open()

out = sys.argv[1]
doc.save(out)
print("wrote", out)
